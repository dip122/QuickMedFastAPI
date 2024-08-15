import os
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from fastapi.middleware.cors import CORSMiddleware
from langchain_google_genai import ChatGoogleGenerativeAI
from fastapi.responses import JSONResponse
from docx import Document
from io import BytesIO
import base64
import uvicorn
# import logging

# logging.basicConfig(level=logging.INFO)

from fastapi import FastAPI
from pydantic import BaseModel

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Update this to match your frontend's URL
    allow_origin_regex="http://localhost:.*",  # Allows all localhost ports
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods
    allow_headers=["*"],  # Allow all headers
)
# Load environment variables
load_dotenv()
os.environ['GOOGLE_API_KEY'] = os.getenv('GOOGLE_API_KEY')
os.environ['SERPER_API_KEY'] = os.getenv('SERPER_API_KEY')


def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations')
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'


# Initialize tools and LLM
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro",
    temperature=0.1,
    max_tokens=5000,
    timeout=None,
)

# Define Agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis",
    backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm,
)

treatment_advisor = Agent(
    role='Treatment Advisor',
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm,
)

# Define Tasks
def find_diagnosis_task(medical_history, symptoms):
    diagnose_task = Task(
        description=(
            f"1. Given that this is a non-critical scenario, analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
            "2. Provide potential conditions that could explain the symptoms.\n"
            "3. Offer advice on steps the user can take at home."
        ),
        expected_output="A preliminary diagnosis with home care advice.",
        agent=diagnostician
    )
    return diagnose_task

def Treatment_task(medical_history, symptoms):
    treatment_task = Task(
        description=(
            "1. Based on the diagnosis, recommend appropriate  basic treatment plans step by step.\n"
            f"1. Recommed some basic treatment plans for those symptoms ({symptoms}).\n"
            f"2. Tell about the probable reasons behind the symptoms.\n"
            f"3. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
            "4. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care.\n"
        ),
        expected_output="A comprehensive treatment plan tailored to the patient's needs.",
        agent=treatment_advisor
    )
    return treatment_task

def create_crew(diagnostician, treatment_advisor, medical_history, symptoms):
    diagnosis_task = find_diagnosis_task(medical_history, symptoms)
    treatment_task = Treatment_task(medical_history, symptoms)
    crew = Crew(
        agents=[diagnostician, treatment_advisor],
        tasks=[diagnosis_task, treatment_task]
    )
    return crew

# Define Pydantic model for the request body
class AIDoctorItem(BaseModel):
    gender: str
    age: int
    medical_history: str
    symptoms: str


# FastAPI POST endpoint
@app.get('/')
def home():
    return "welcome to health checker application fast api"
    
@app.post('/diagnosis')
async def diagnosis_medical_agent(item: AIDoctorItem):
    # Extract data from request
    gender = item.gender
    age = item.age
    medical_history = item.medical_history
    symptoms = item.symptoms
    
    # Create and execute crew
    crew = create_crew(diagnostician, treatment_advisor, medical_history, symptoms)
    result = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
    print(result.raw)
    # Process result (assuming result is a string; modify if different)
    return result

@app.get("/test")
async def test():
    return {"message": "This is a test endpoint"}


if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=5000)